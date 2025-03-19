import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document


def remove_images_from_element(element):
    """递归删除文档元素（段落、表格等）中的图片"""
    for para in element.paragraphs:
        for run in para.runs:
            # 查找图片对应的 drawing 元素
            drawings = run._element.xpath('.//w:drawing')
            for drawing in drawings:
                drawing.getparent().remove(drawing)
    for table in element.tables:
        for row in table.rows:
            for cell in row.cells:
                remove_images_from_element(cell)


def remove_images_from_docx(input_path, output_path):
    """加载文档，删除所有图片，并保存为新的文档"""
    doc = Document(input_path)
    # 处理正文
    remove_images_from_element(doc)
    # 处理页眉和页脚
    for section in doc.sections:
        remove_images_from_element(section.header)
        remove_images_from_element(section.footer)
    doc.save(output_path)


def select_input_file():
    """弹出对话框选择输入文件"""
    return filedialog.askopenfilename(title="选择要处理的 Word 文件", filetypes=[("Word文档", "*.docx")])


def select_output_file():
    """弹出对话框选择输出文件路径"""
    return filedialog.asksaveasfilename(title="保存处理后的 Word 文件", defaultextension=".docx",
                                        filetypes=[("Word文档", "*.docx")])


def main():
    # 初始化 Tkinter 主窗口并隐藏
    root = tk.Tk()
    root.withdraw()

    input_file = select_input_file()
    if not input_file:
        messagebox.showerror("错误", "未选择输入文件")
        return

    output_file = select_output_file()
    if not output_file:
        messagebox.showerror("错误", "未选择输出文件路径")
        return

    try:
        remove_images_from_docx(input_file, output_file)
        messagebox.showinfo("完成", f"处理完成，已保存到\n{output_file}")
    except Exception as e:
        messagebox.showerror("错误", f"处理过程中出错：\n{e}")


if __name__ == '__main__':
    main()
